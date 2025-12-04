"""
HWP 텍스트 변환 어댑터

여러 HWP 변환 방법을 제공하고 환경에 따라 자동 선택합니다.
"""
import logging
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, List
from docx import Document  

logger = logging.getLogger(__name__)

# ========================================
# 내부 유틸리티: soffice 경로 탐색 (Windows 대응)
# ========================================

def _get_soffice_cmd() -> str:
    """
    OS에 따라 LibreOffice CLI(soffice) 실행 파일 경로 결정.

    - Windows:
        기본 설치 경로(C:\\Program Files\\LibreOffice\\program\\soffice.exe 등)를 우선 탐색
        없으면 마지막으로 'soffice' (PATH에 등록된 경우) 사용.
    - Linux/Mac:
        그냥 'soffice' 사용 (PATH에 있다고 가정).
    """
    if os.name == "nt":  # Windows
        candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for c in candidates:
            if Path(c).exists():
                return c
        # 그래도 못 찾으면 PATH에 있는 soffice에 마지막으로 기대
        return "soffice"
    else:
        # Linux / Mac 은 PATH 에서 찾도록 위임
        return "soffice"


# ========================================
# HWP 변환 방법 체크
# ========================================

def _check_hwp5txt_available() -> bool:
    """hwp5txt CLI 도구 사용 가능 여부 확인"""
    try:
        result = subprocess.run(
            ["hwp5txt", "--version"],
            capture_output=True,
            timeout=5
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def _check_libreoffice_available() -> bool:
    """LibreOffice CLI 도구 사용 가능 여부 확인"""
    try:
        cmd = _get_soffice_cmd()

        # 1) Windows에서는 exe 파일이 실제로 존재하면 그냥 사용 가능하다고 판단
        if os.name == "nt":
            # cmd가 "C:\Program Files\LibreOffice\program\soffice.exe" 같은 절대경로일 수도 있고
            # 그냥 "soffice"일 수도 있으니, 절대경로인 경우만 파일 체크
            p = Path(cmd)
            if p.is_absolute() and p.exists():
                logger.info(f"[LibreOffice] Detected soffice at {p}")
                return True

        # 2) 그 외(또는 절대경로가 아닌 경우)에는 --version 호출해서 확인
        result = subprocess.run(
            [cmd, "--version"],
            capture_output=True,
            timeout=10,
            text=True,
        )
        return result.returncode == 0

    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


# ========================================
# 변환 방법 1: hwp5txt (세희 방식, Linux 권장)
# ========================================

def convert_hwp_with_hwp5txt(hwp_path: str) -> str:
    """
    hwp5txt CLI 도구로 HWP → 텍스트 변환

    ⚠️ 세희 코드에서 가져온 방식 (prompt.txt:65-76)
    ⚠️ Linux/Mac 전용 (hwp5 패키지 필요)

    설치:
        - Ubuntu: sudo apt-get install hwp5 && pip install hwp5
        - Mac: brew install hwp5 && pip install hwp5

    Args:
        hwp_path: HWP 파일 경로

    Returns:
        str: 추출된 텍스트

    Raises:
        FileNotFoundError: hwp5txt 명령어가 없을 때
        subprocess.CalledProcessError: 변환 실패 시
    """
    hwp_path = Path(hwp_path).resolve()

    if not hwp_path.exists():
        raise FileNotFoundError(f"HWP file not found: {hwp_path}")

    logger.info(f"[hwp5txt] Converting HWP: {hwp_path.name}")

    try:
        result = subprocess.run(
            ["hwp5txt", str(hwp_path)],
            capture_output=True,
            text=True,
            check=True,
            timeout=60  # 60초 타임아웃
        )

        text = result.stdout
        logger.info(f"[hwp5txt] Extracted {len(text)} chars from {hwp_path.name}")
        return text

    except FileNotFoundError:
        logger.error("hwp5txt not found. Install: pip install hwp5 (Linux/Mac only)")
        raise

    except subprocess.CalledProcessError as e:
        logger.error(f"hwp5txt conversion failed: {e.stderr}")
        raise

    except subprocess.TimeoutExpired:
        logger.error(f"hwp5txt timeout (>60s) for {hwp_path.name}")
        raise

def _extract_text_from_docx(docx_path: Path) -> str:
    """DOCX 파일에서 텍스트 추출 (단락 기준)"""
    doc = Document(str(docx_path))
    paragraphs: list[str] = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            paragraphs.append(t)

    # 필요하면 테이블 텍스트도 추가
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n".join(paragraphs)


# ========================================
# 변환 방법 2: LibreOffice CLI (크로스 플랫폼)
# ========================================

def convert_hwp_with_libreoffice(hwp_path: str) -> str:
    """
    LibreOffice로 HWP → DOCX → 텍스트 변환

    기존: HWP → PDF → 텍스트
    변경: HWP → DOCX → python-docx 로 텍스트 추출
    (DOCX 쪽이 줄바꿈/한글 인코딩이 더 안정적이라 이 경로 권장)
    """
    hwp_path = Path(hwp_path).resolve()

    if not hwp_path.exists():
        raise FileNotFoundError(f"HWP file not found: {hwp_path}")

    cmd = _get_soffice_cmd()
    logger.info(f"[LibreOffice] Converting HWP to DOCX: {hwp_path.name} (cmd={cmd})")

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)

        try:
            # 1) HWP → DOCX 변환
            result = subprocess.run(
                [
                    cmd,
                    "--headless",
                    "--infilter=Hwp2002_File",
                    # ★ 여기가 핵심: pdf가 아니라 docx로 변환
                    "--convert-to", "docx:MS Word 2007 XML",
                    str(hwp_path),
                    "--outdir", str(tmpdir_path),
                ],
                capture_output=True,
                text=True,
                timeout=120,
            )

            if result.returncode != 0:
                logger.error(f"[LibreOffice] stderr: {result.stderr}")

            docx_file = tmpdir_path / f"{hwp_path.stem}.docx"
            if not docx_file.exists():
                logger.error("[LibreOffice] DOCX not created from HWP")
                raise RuntimeError(f"Conversion failed: {result.stderr}")

            logger.info(
                f"[LibreOffice] DOCX created: {docx_file} "
                f"(size={docx_file.stat().st_size} bytes)"
            )

            # 2) DOCX에서 텍스트 추출
            full_text = _extract_text_from_docx(docx_file)
            logger.info(f"[LibreOffice] Extracted {len(full_text)} chars via DOCX")
            return full_text

        except FileNotFoundError:
            logger.error("LibreOffice (soffice) not found. Install LibreOffice first.")
            raise

        except subprocess.TimeoutExpired:
            logger.error(f"LibreOffice timeout (>120s) for {hwp_path.name}")
            raise RuntimeError("LibreOffice conversion timeout")


# ========================================
# 변환 방법 3: pyhwp (Python 2 전용, Deprecated)
# ========================================

def convert_hwp_with_pyhwp(hwp_path: str) -> str:
    """
    pyhwp 라이브러리로 HWP → 텍스트 변환

    ❌ Python 2 전용 (Python 3에서 설치 실패)
    ⚠️ Deprecated: hwp5txt 또는 LibreOffice 사용 권장

    Args:
        hwp_path: HWP 파일 경로

    Returns:
        str: 추출된 텍스트

    Raises:
        ImportError: pyhwp 설치 안되어 있을 때
    """
    try:
        import pyhwp
    except ImportError:
        logger.error("pyhwp not installed (Python 2 only)")
        raise ImportError("pyhwp requires Python 2. Use hwp5txt or LibreOffice instead.")

    hwp_path = Path(hwp_path).resolve()

    if not hwp_path.exists():
        raise FileNotFoundError(f"HWP file not found: {hwp_path}")

    logger.info(f"[pyhwp] Converting HWP: {hwp_path.name}")

    text = ""
    try:
        doc = pyhwp.HWPDocument(str(hwp_path))
        for para in doc.bodytext.paragraphs:
            for run in para.text:
                text += run.text
            text += "\n"

        logger.info(f"[pyhwp] Extracted {len(text)} chars from {hwp_path.name}")
        return text

    except Exception as e:
        logger.error(f"pyhwp conversion failed: {e}")
        raise RuntimeError(f"pyhwp conversion failed: {e}")


# ========================================
# 통합 변환 함수 (자동 선택)
# ========================================

def convert_hwp_to_text(hwp_path: str, method: Optional[str] = None) -> str:
    """
    HWP 파일을 텍스트로 변환 (여러 방법 중 자동 선택)

    우선순위:
    1. method 파라미터 지정 시: 해당 방법 사용
    2. hwp5txt 사용 가능 시: hwp5txt (세희 방식, 가장 안정적)
    3. LibreOffice 사용 가능 시: LibreOffice (크로스 플랫폼)
    4. 모두 실패 시: 빈 문자열 반환 (graceful fallback)

    Args:
        hwp_path: HWP 파일 경로
        method: 변환 방법 ("hwp5txt" | "libreoffice" | "pyhwp" | None)

    Returns:
        str: 추출된 텍스트 (실패 시 빈 문자열)
    """
    hwp_path = Path(hwp_path).resolve()

    if not hwp_path.exists():
        logger.error(f"HWP file not found: {hwp_path}")
        return ""

    # 1. 지정된 방법 사용
    if method:
        logger.info(f"Using specified method: {method}")
        try:
            if method == "hwp5txt":
                return convert_hwp_with_hwp5txt(str(hwp_path))
            elif method == "libreoffice":
                return convert_hwp_with_libreoffice(str(hwp_path))
            elif method == "pyhwp":
                return convert_hwp_with_pyhwp(str(hwp_path))
            else:
                logger.warning(f"Unknown method: {method}, falling back to auto")
        except Exception as e:
            logger.error(f"Specified method '{method}' failed: {e}")
            return ""

    # 2. hwp5txt 시도 (세희 방식, 최우선)
    if _check_hwp5txt_available():
        logger.info("Using hwp5txt (preferred method)")
        try:
            return convert_hwp_with_hwp5txt(str(hwp_path))
        except Exception as e:
            logger.warning(f"hwp5txt failed: {e}, trying next method")

    # 3. LibreOffice 시도
    if _check_libreoffice_available():
        logger.info("Using LibreOffice")
        try:
            return convert_hwp_with_libreoffice(str(hwp_path))
        except Exception as e:
            logger.warning(f"LibreOffice failed: {e}")

    # 4. pyhwp 시도 (Deprecated)
    try:
        logger.info("Using pyhwp (deprecated)")
        return convert_hwp_with_pyhwp(str(hwp_path))
    except Exception as e:
        logger.warning(f"pyhwp failed: {e}")

    # 5. 모두 실패 - graceful fallback
    logger.error(f"All HWP conversion methods failed for {hwp_path.name}")
    logger.error("Please install one of: hwp5txt (Linux), LibreOffice (any OS)")
    return ""


# ========================================
# 유틸리티 함수
# ========================================

def get_available_methods() -> List[str]:
    """사용 가능한 HWP 변환 방법 목록 반환"""
    methods: List[str] = []

    if _check_hwp5txt_available():
        methods.append("hwp5txt")

    if _check_libreoffice_available():
        methods.append("libreoffice")

    try:
        import pyhwp  # noqa: F401
        methods.append("pyhwp")
    except ImportError:
        pass

    return methods


def get_recommended_method() -> Optional[str]:
    """환경에 맞는 권장 HWP 변환 방법 반환"""
    if _check_hwp5txt_available():
        return "hwp5txt"  # 세희 방식, 가장 안정적
    elif _check_libreoffice_available():
        return "libreoffice"  # 크로스 플랫폼
    else:
        return None


if __name__ == "__main__":
    # 사용 가능한 변환 방법 출력
    print("Available HWP conversion methods:")
    methods = get_available_methods()
    if methods:
        for method in methods:
            print(f"  - {method}")
    else:
        print("  (none)")

    recommended = get_recommended_method()
    if recommended:
        print(f"\nRecommended: {recommended}")
    else:
        print("\n⚠️ No HWP converter available. Install hwp5 or LibreOffice.")
